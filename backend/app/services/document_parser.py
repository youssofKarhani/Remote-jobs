"""CV Document Parser supporting Plain Text, PDF, and DOCX formats."""

import io
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Optional

import docx
from pypdf import PdfReader


class DocumentParsingError(ValueError):
    """Base exception for document parsing failures."""
    pass


class EmptyDocumentError(DocumentParsingError):
    """Raised when an uploaded document contains no extractable text."""
    pass


class UnsupportedFileFormatError(DocumentParsingError):
    """Raised when the document file extension/MIME type is not supported."""
    pass


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF binary using pypdf."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        extracted_pages = []
        for page_idx, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                extracted_pages.append(page_text.strip())
        return "\n\n".join(extracted_pages)
    except Exception as e:
        raise DocumentParsingError(f"Failed to extract text from PDF: {str(e)}") from e


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX binary using python-docx with zipfile fallback."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))
        if parts:
            return "\n\n".join(parts)
    except Exception:
        # Fallback to direct XML parsing of Word OpenXML archive
        pass

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            xml_content = zf.read("word/document.xml")
            tree = ET.fromstring(xml_content)
            # Namespace for Word OpenXML
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            text_nodes = tree.findall(".//w:t", ns)
            text_list = [node.text for node in text_nodes if node.text]
            return " ".join(text_list)
    except Exception as e:
        raise DocumentParsingError(f"Failed to extract text from DOCX: {str(e)}") from e


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text binary trying common encodings."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]
    for enc in encodings:
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def parse_document(file_bytes: bytes, filename: str) -> str:
    """Parse binary document and return sanitized plaintext string.
    
    Supports .pdf, .docx, .doc, .txt, .md.
    Raises EmptyDocumentError if no readable text is found.
    Raises UnsupportedFileFormatError if extension is unsupported.
    """
    if not file_bytes:
        raise EmptyDocumentError("The uploaded file is empty (0 bytes).")

    ext = filename.lower().split(".")[-1] if "." in filename else ""

    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        text = extract_text_from_docx(file_bytes)
    elif ext in ("txt", "md", "text", "rtf", ""):
        text = extract_text_from_txt(file_bytes)
    else:
        # Attempt fallback to txt decoding before erroring
        try:
            text = extract_text_from_txt(file_bytes)
        except Exception:
            raise UnsupportedFileFormatError(
                f"Unsupported file format '.{ext}'. Please upload a PDF, DOCX, or TXT resume."
            )

    # Sanitize and normalize whitespace
    cleaned_text = re.sub(r"\r\n|\r", "\n", text).strip()
    if not cleaned_text or len(cleaned_text) < 10:
        raise EmptyDocumentError(
            "The document contains no readable text or is a scanned image. "
            "Please provide a text-based PDF, DOCX, or paste plain text directly."
        )

    return cleaned_text
